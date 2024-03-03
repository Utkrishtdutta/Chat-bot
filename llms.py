from transformers import pipeline
from PyPDF2 import PdfReader
from docx import Document

def read_file(file):
    file_content = ''
    if file.type == 'text/plain':
        file_content += file.read().decode('utf-8')
    elif file.type == 'application/pdf':
        pdf_reader = PdfReader(file)
        for page_num in range(len(pdf_reader.pages)):
            file_content += pdf_reader.pages[page_num].extract_text()
    elif file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = Document(file)
        for paragraph in doc.paragraphs:
            file_content += paragraph.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    file_content += cell.text
    return file_content

def generate_gemma(input_text):
    pipe = pipeline("text-generation", model="google/gemma-7b",max_length=100)
    response = pipe(input_text)
    return response[0]['generated_text']


def generate_openai(input_text):
    pipe = pipeline("text-generation", model="openai-community/gpt2",max_length=100)
    response = pipe(input_text)
    return response[0]['generated_text']